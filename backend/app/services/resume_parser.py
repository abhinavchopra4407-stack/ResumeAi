import re
import fitz
import pdfplumber
from docx import Document
from typing import Dict, List, Optional, Any
import json
from datetime import datetime
import logging
import io

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        # Email pattern
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        
        # Phone pattern - handles international and various formats
        self.phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,4}[-.\s]?\d{2,4}'
        
        # Common skills database
        self.skills_db = {
            'programming': [
                'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Ruby', 'Go', 'Rust',
                'Swift', 'Kotlin', 'PHP', 'Perl', 'R', 'MATLAB', 'Scala', 'Groovy'
            ],
            'web': [
                'React', 'Angular', 'Vue.js', 'Node.js', 'Express.js', 'Django', 'Flask', 'FastAPI',
                'Spring Boot', 'ASP.NET', 'Laravel', 'Ruby on Rails', 'Next.js', 'Nuxt.js',
                'HTML5', 'CSS3', 'SASS', 'LESS', 'Bootstrap', 'Tailwind CSS', 'Material-UI'
            ],
            'database': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'SQLite', 'Redis', 'Elasticsearch', 'Cassandra',
                'DynamoDB', 'Oracle', 'SQL Server', 'Firebase', 'Neo4j'
            ],
            'cloud': [
                'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Terraform', 'Ansible', 'Jenkins',
                'GitLab CI', 'GitHub Actions', 'CircleCI', 'Travis CI', 'Prometheus', 'Grafana'
            ],
            'data_science': [
                'Machine Learning', 'Deep Learning', 'NLP', 'Computer Vision', 'Data Science',
                'Data Analytics', 'Statistics', 'TensorFlow', 'PyTorch', 'Scikit-learn',
                'Pandas', 'NumPy', 'SciPy', 'Matplotlib', 'Seaborn', 'Tableau', 'Power BI',
                'Hadoop', 'Spark', 'Kafka', 'Airflow'
            ],
            'others': [
                'Git', 'Agile', 'Scrum', 'Kanban', 'JIRA', 'Confluence', 'Leadership',
                'Project Management', 'Team Management', 'Communication', 'Problem Solving',
                'Critical Thinking', 'Time Management', 'Team Collaboration'
            ]
        }
        
        # Section detection patterns
        self.section_patterns = {
            'contact': [
                r'(?i)^\s*(contact|personal\s+info|contact\s+information|phone|email)\s*$'
            ],
            'summary': [
                r'(?i)^\s*(summary|professional\s+summary|profile|objective|about\s+me|candidate\s+summary)\s*$'
            ],
            'experience': [
                r'(?i)^\s*(work\s+experience|professional\s+experience|employment|experience|work\s+history|career|employment\s+history)\s*$'
            ],
            'education': [
                r'(?i)^\s*(education|academic\s+background|qualifications|degrees|educational\s+qualifications)\s*$'
            ],
            'skills': [
                r'(?i)^\s*(skills|technical\s+skills|core\s+competencies|expertise|technologies|tech\s+skills)\s*$'
            ],
            'projects': [
                r'(?i)^\s*(projects|personal\s+projects|academic\s+projects|key\s+projects|portfolio)\s*$'
            ],
            'certifications': [
                r'(?i)^\s*(certifications|certificates|licenses|professional\s+certifications)\s*$'
            ],
            'languages': [
                r'(?i)^\s*(languages|spoken\s+languages|language\s+proficiency)\s*$'
            ],
            'achievements': [
                r'(?i)^\s*(achievements|awards|honors|recognitions|publications)\s*$'
            ]
        }

    @classmethod
    def parse_text(cls, text: str) -> Dict:
        """
        Parse resume from text
        """
        parser = cls()
        return parser._extract_entities(text)

    async def parse(self, file_content: bytes, file_name: str) -> Dict:
        """
        Parse resume from file content
        """
        try:
            file_extension = file_name.split('.')[-1].lower()
            
            # Extract text based on file type
            if file_extension == 'pdf':
                text = self._extract_pdf(file_content)
            elif file_extension == 'docx':
                text = self._extract_docx(file_content)
            else:
                # If it's a text file or unknown, try to decode as text
                try:
                    text = file_content.decode('utf-8', errors='ignore')
                except:
                    raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Log the extracted text length for debugging
            logger.info(f"Extracted text length: {len(text)} characters from {file_name}")
            
            if not text or len(text.strip()) == 0:
                logger.warning(f"No text extracted from {file_name}")
                return self._get_empty_parsed_data()
            
            # Parse text into structured data
            return self._extract_entities(text)
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_name}: {str(e)}", exc_info=True)
            raise

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Method 1: Try PyMuPDF (faster, more reliable)
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
            doc.close()
            
            if text.strip():
                logger.info(f"PyMuPDF extracted {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {str(e)}")
        
        # Method 2: Try pdfplumber
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"pdfplumber extracted {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 3: Try PyMuPDF with different settings (if available)
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                page_text = page.get_text("dict")
                if page_text:
                    # Extract text from dictionary
                    for block in page_text.get('blocks', []):
                        if 'lines' in block:
                            for line in block['lines']:
                                for span in line.get('spans', []):
                                    if span.get('text'):
                                        text += span['text'] + " "
                    text += "\n"
            doc.close()
            
            if text.strip():
                logger.info(f"PyMuPDF (dict) extracted {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF dict extraction failed: {str(e)}")
        
        return text

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            logger.info(f"DOCX extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    def _get_empty_parsed_data(self) -> Dict:
        """Return empty parsed data structure"""
        return {
            "name": "Unknown Candidate",
            "email": "",
            "phone": "",
            "skills": [],
            "experience": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "languages": [],
            "achievements": [],
            "summary": "",
            "full_text": "",
            "word_count": 0,
            "line_count": 0,
            "sections": {}
        }

    def _extract_entities(self, text: str) -> Dict:
        """
        Extract structured information from resume text
        """
        if not text or len(text.strip()) == 0:
            return self._get_empty_parsed_data()
        
        # Split into lines, keeping empty ones for structure
        lines = [line.rstrip() for line in text.split('\n')]
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        # Extract basic information
        name = self._extract_name(non_empty_lines, text)
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        
        # Identify sections
        sections = self._identify_sections(lines)
        
        # Extract content from each section
        skills = self._extract_skills(sections.get('skills', []))
        experience = self._extract_experience(sections.get('experience', []))
        education = self._extract_education(sections.get('education', []))
        projects = self._extract_projects(sections.get('projects', []))
        certifications = self._extract_certifications(sections.get('certifications', []))
        languages = self._extract_languages(sections.get('languages', []))
        achievements = self._extract_achievements(sections.get('achievements', []))
        summary = self._extract_summary(sections.get('summary', []))
        
        # If no sections were identified, try to parse as a simple text resume
        if not any([skills, experience, education, projects]):
            logger.info("No sections identified, trying simple text parsing")
            simple_parsed = self._parse_simple_resume(text)
            return simple_parsed
        
        return {
            "name": name if name else "Unknown Candidate",
            "email": email if email else "",
            "phone": phone if phone else "",
            "skills": skills,
            "experience": experience,
            "education": education,
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
            "achievements": achievements,
            "summary": summary,
            "full_text": text,
            "word_count": len(text.split()),
            "line_count": len(non_empty_lines),
            "sections": sections
        }

    def _parse_simple_resume(self, text: str) -> Dict:
        """Parse a simple resume without clear sections"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Try to find experience lines (containing dates)
        experience = []
        edu_lines = []
        project_lines = []
        
        date_pattern = r'(19|20)\d{2}'
        
        for line in lines:
            if re.search(date_pattern, line):
                if any(keyword in line.lower() for keyword in ['bachelor', 'master', 'phd', 'university', 'college']):
                    edu_lines.append(line)
                elif any(keyword in line.lower() for keyword in ['project', 'built', 'developed', 'created']):
                    project_lines.append(line)
                else:
                    experience.append(line)
            elif 'project' in line.lower() or 'developed' in line.lower():
                project_lines.append(line)
            elif any(keyword in line.lower() for keyword in ['bachelor', 'master', 'phd', 'university', 'college']):
                edu_lines.append(line)
            elif len(line) > 20:
                experience.append(line)
        
        # Extract skills from text
        skills = self._extract_skills_from_text(text)
        
        return {
            "name": self._extract_name(lines, text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "skills": skills,
            "experience": self._extract_experience(experience[:10]),
            "education": self._extract_education(edu_lines[:3]),
            "projects": self._extract_projects(project_lines[:3]),
            "certifications": [],
            "languages": [],
            "achievements": [],
            "summary": self._extract_summary(lines[:5]),
            "full_text": text,
            "word_count": len(text.split()),
            "line_count": len(lines),
            "sections": {}
        }

    def _extract_skills_from_text(self, text: str) -> List[str]:
        """Extract skills directly from text"""
        found_skills = []
        text_lower = text.lower()
        
        for category, skills in self.skills_db.items():
            for skill in skills:
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.append(skill)
        
        # Remove duplicates while preserving order
        unique_skills = []
        for skill in found_skills:
            if skill not in unique_skills:
                unique_skills.append(skill)
        
        return unique_skills

    def _extract_name(self, lines: List[str], text: str) -> str:
        """Extract candidate name"""
        # Try to find name in first few lines
        for i in range(min(10, len(lines))):
            line = lines[i]
            # Name is usually: 2-3 words, no special characters, not a section header
            words = line.split()
            if 2 <= len(words) <= 4 and not any(c.isdigit() for c in line):
                # Check if it's a section header
                is_section = False
                for patterns in self.section_patterns.values():
                    if any(re.search(pattern, line, re.IGNORECASE) for pattern in patterns):
                        is_section = True
                        break
                
                if not is_section and len(line) < 50:
                    # Check if it contains common header keywords
                    if not any(keyword in line.lower() for keyword in 
                              ['resume', 'cv', 'curriculum', 'vitae', 'email', 'phone', 'contact']):
                        # Properly format name
                        name_parts = line.split()
                        formatted_name = ' '.join([part.title() for part in name_parts])
                        return formatted_name
        
        # Fallback: look for name format patterns in first 1000 chars
        first_chars = text[:1000]
        name_patterns = [
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$',  # John Smith
            r'^[A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+$',  # John D. Smith
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+,\s+[A-Z][a-z]+$',  # Smith, John
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, first_chars, re.MULTILINE)
            if match:
                return match.group(0)
        
        # If we have a name in the first line that's not too long
        if lines and 2 <= len(lines[0].split()) <= 4:
            return lines[0]
        
        return "Unknown Candidate"

    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        match = re.search(self.email_pattern, text)
        return match.group(0) if match else ""

    def _extract_phone(self, text: str) -> str:
        """Extract phone number"""
        # Try different phone patterns
        phone_patterns = [
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,4}[-.\s]?\d{2,4}',
            r'\d{3}[-.]\d{3}[-.]\d{4}',
            r'\(\d{3}\)\s*\d{3}[-.]\d{4}',
            r'\d{10}'
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        
        return ""

    def _identify_sections(self, lines: List[str]) -> Dict[str, List[str]]:
        """
        Identify and extract sections from resume
        """
        sections = {key: [] for key in self.section_patterns.keys()}
        current_section = 'summary'
        
        # Look for section headers
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if this line is a section header
            found_section = None
            line_lower = line.lower()
            
            for section_name, patterns in self.section_patterns.items():
                if any(re.search(pattern, line_lower) for pattern in patterns):
                    # Check if it's likely a header (short line, often all caps or with colons)
                    if len(line) < 40 and (line.isupper() or ':' in line or line_lower in ['skills', 'experience', 'education']):
                        found_section = section_name
                        break
            
            if found_section:
                current_section = found_section
                # Skip section header line
                i += 1
                continue
            
            # Add content to current section
            if current_section in sections:
                sections[current_section].append(line)
            
            i += 1
        
        # Clean up sections - remove empty or very short entries
        for section_name in sections:
            # Filter out lines that are likely section headers
            filtered_lines = []
            for line in sections[section_name]:
                is_header = False
                line_lower = line.lower()
                for patterns in self.section_patterns.values():
                    if any(re.search(pattern, line_lower) for pattern in patterns):
                        if len(line) < 40:
                            is_header = True
                            break
                if not is_header and len(line.strip()) > 2:
                    filtered_lines.append(line)
            sections[section_name] = filtered_lines
        
        return sections

    def _extract_skills(self, skill_lines: List[str]) -> List[str]:
        """Extract skills from text"""
        if not skill_lines:
            return []
        
        text = ' '.join(skill_lines).lower()
        found_skills = []
        
        # Check each skill category
        for category, skills in self.skills_db.items():
            for skill in skills:
                # Match skill as whole word or phrase
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text):
                    found_skills.append(skill)
        
        # If no skills found with exact matching, try fuzzy matching
        if not found_skills:
            # Check for skills that appear as substrings
            for category, skills in self.skills_db.items():
                for skill in skills:
                    if skill.lower() in text:
                        found_skills.append(skill)
        
        # Remove duplicates while preserving order
        unique_skills = []
        for skill in found_skills:
            if skill not in unique_skills:
                unique_skills.append(skill)
        
        return unique_skills

    def _extract_experience(self, exp_lines: List[str]) -> List[Dict]:
        """Extract work experience entries"""
        experiences = []
        if not exp_lines:
            return experiences
        
        # Try to identify experience blocks
        current_exp = {}
        date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}'
        year_pattern = r'(19|20)\d{2}\s*[-–]\s*(?:19|20)\d{2}|(?:19|20)\d{2}'
        
        for line in exp_lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for date patterns
            has_date = re.search(date_pattern, line, re.IGNORECASE) or re.search(year_pattern, line)
            
            # If line looks like a job title or has date
            if has_date or (len(line.split()) < 8 and any(c.isupper() for c in line[:20])):
                if current_exp and current_exp.get('title'):
                    experiences.append(current_exp)
                current_exp = {
                    "title": line,
                    "company": "",
                    "duration": "",
                    "description": []
                }
            elif current_exp:
                # If line looks like a company name (short, proper case)
                if not current_exp['company'] and len(line.split()) < 5 and any(c.isupper() for c in line):
                    current_exp['company'] = line
                elif len(line) > 10:
                    current_exp['description'].append(line)
        
        # Add the last experience
        if current_exp and current_exp.get('title'):
            experiences.append(current_exp)
        
        # Clean up and format
        for exp in experiences:
            if not exp.get('company') and exp.get('description'):
                # First description line might be company
                first_desc = exp['description'][0] if exp['description'] else ''
                if len(first_desc.split()) < 5:
                    exp['company'] = first_desc
                    exp['description'] = exp['description'][1:]
            
            if exp.get('description'):
                exp['description'] = '\n'.join(exp['description'])
            else:
                exp['description'] = ''
        
        return experiences[:5]

    def _extract_education(self, edu_lines: List[str]) -> List[Dict]:
        """Extract education entries"""
        education = []
        if not edu_lines:
            return education
        
        current_edu = {}
        degree_keywords = ['bachelor', 'master', 'phd', 'b.sc', 'm.sc', 'mba', 'b.tech', 'm.tech', 
                          'b.e', 'm.e', 'ba', 'ma', 'ms', 'bs', 'associate', 'doctorate']
        
        for line in edu_lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains degree information
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in degree_keywords):
                if current_edu and current_edu.get('degree'):
                    education.append(current_edu)
                current_edu = {
                    "degree": line,
                    "institution": "",
                    "year": "",
                    "grade": ""
                }
            elif current_edu and not current_edu.get('institution'):
                # Likely institution name (short line, proper case)
                if len(line) > 5 and len(line) < 60:
                    current_edu['institution'] = line
                else:
                    # Maybe it's part of description
                    if 'description' not in current_edu:
                        current_edu['description'] = []
                    current_edu['description'].append(line)
            elif current_edu:
                # Add to description if it's additional info
                if 'description' not in current_edu:
                    current_edu['description'] = []
                current_edu['description'].append(line)
        
        if current_edu and current_edu.get('degree'):
            education.append(current_edu)
        
        # Clean up
        for edu in education:
            if 'description' in edu and edu['description']:
                edu['description'] = '\n'.join(edu['description'])
            else:
                edu['description'] = ''
        
        return education[:3]

    def _extract_projects(self, project_lines: List[str]) -> List[Dict]:
        """Extract project entries"""
        projects = []
        if not project_lines:
            return projects
        
        current_project = {}
        for line in project_lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this is a project title (usually shorter, may have technologies)
            if len(line.split()) < 10 and not current_project.get('name'):
                current_project = {
                    "name": line,
                    "description": [],
                    "technologies": []
                }
            elif current_project:
                current_project['description'].append(line)
            else:
                # If no current project, start one
                current_project = {
                    "name": f"Project {len(projects) + 1}",
                    "description": [line],
                    "technologies": []
                }
        
        if current_project and current_project.get('description'):
            # Look for technologies in description
            all_text = ' '.join(current_project['description'])
            tech_found = []
            for category, skills in self.skills_db.items():
                for skill in skills:
                    if skill.lower() in all_text.lower():
                        tech_found.append(skill)
            current_project['technologies'] = tech_found[:5]
            current_project['description'] = '\n'.join(current_project['description'])
            projects.append(current_project)
        
        return projects

    def _extract_certifications(self, cert_lines: List[str]) -> List[str]:
        """Extract certifications"""
        certifications = []
        if not cert_lines:
            return certifications
        
        cert_keywords = ['certified', 'certificate', 'certification', 'aws', 'azure', 'google cloud',
                        'scrum', 'project management', 'pmp', 'itil', 'ccna', 'ccnp', 'ceh', 'cissp']
        
        for line in cert_lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in cert_keywords):
                certifications.append(line)
            elif len(line) < 50 and (line.isupper() or ':' in line):
                # May be a certification line
                certifications.append(line)
        
        return certifications[:5]

    def _extract_languages(self, lang_lines: List[str]) -> List[str]:
        """Extract languages"""
        languages = []
        if not lang_lines:
            return languages
        
        lang_db = ['English', 'Spanish', 'French', 'German', 'Chinese', 'Japanese', 'Korean',
                  'Portuguese', 'Italian', 'Dutch', 'Russian', 'Arabic', 'Hindi', 'Bengali']
        
        text = ' '.join(lang_lines).lower()
        
        for lang in lang_db:
            if lang.lower() in text:
                languages.append(lang)
        
        return languages

    def _extract_achievements(self, achievement_lines: List[str]) -> List[str]:
        """Extract achievements"""
        achievements = []
        if not achievement_lines:
            return achievements
        
        for line in achievement_lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['award', 'achievement', 'recognition', 'honor']):
                achievements.append(line)
            elif len(line) > 30 and any(keyword in line.lower() for keyword in ['won', 'received', 'ranked']):
                achievements.append(line)
        
        return achievements[:5]

    def _extract_summary(self, summary_lines: List[str]) -> str:
        """Extract professional summary"""
        if not summary_lines:
            return ""
        
        summary = ' '.join(summary_lines[:5])  # First few lines of summary
        return summary[:500]